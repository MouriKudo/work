"""严格按 8 项任务交付清单生成 Word 报告与 ZIP 压缩包。"""

import os
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT = Path(__file__).resolve().parent.parent


# ───────────────── 工具函数 ─────────────────

def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True)
        tc = table.rows[0].cells[j]._element.get_or_add_tcPr()
        tc.append(tc.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "4472C4"}))
        for r in table.rows[0].cells[j].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            a = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(table.rows[i + 1].cells[j], val, size=9, align=a)
        if i % 2 == 1:
            for j in range(len(headers)):
                tc = table.rows[i + 1].cells[j]._element.get_or_add_tcPr()
                tc.append(tc.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "D9E2F3"}))


def add_fig(doc, rel_path, caption, width=Inches(5.8)):
    full = PROJECT / rel_path
    if not full.exists():
        doc.add_paragraph(f"[图片缺失：{rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(full), width=width)
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(80, 80, 80)


def deliverable_list(doc, items):
    """打印交付物清单，每项前面带 ✔ 标记。"""
    for item in items:
        p = doc.add_paragraph()
        r = p.add_run("✔  ")
        r.bold = True
        r.font.color.rgb = RGBColor(0, 128, 0)
        r.font.size = Pt(10)
        r2 = p.add_run(item)
        r2.font.size = Pt(10)


# ───────────────── 主报告 ─────────────────

def create_report(output_path: Path):
    doc = Document()

    # 全局样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25

    # ═══ 封面 ═══
    title = doc.add_heading("LUNA16 PBIP-Lite 肺结节候选分类", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("增量实验成果报告")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(68, 114, 196)
    doc.add_paragraph("")
    doc.add_paragraph("")

    # ═══ 任务总览 ═══
    doc.add_heading("任务总览", level=1)
    overview = [
        "任务 1：结果汇总 + K 值补充实验",
        "任务 2：消融实验（组件 / K 值 / β 值）",
        "任务 3：图像退化算子实现",
        "任务 4：退化鲁棒性评估",
        "任务 5：Robust-PBIP 训练",
        "任务 6：Grad-CAM + 四类样本可视化",
        "任务 7：Top-3 相似病例检索",
        "任务 8：LIDC-IDRI 外部验证可行性分析",
    ]
    for item in overview:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 1：主结果表、指标对比图、K值结果、方法结论
    # ══════════════════════════════════════════
    doc.add_heading("任务 1　结果汇总 + K 值补充实验", level=1)
    doc.add_paragraph(
        "统一汇总 4 种方法（ResNet18 基线 / ResNet18 增强 / PBIP-Lite 无对比损失 / PBIP-Lite 完整）"
        "在 seed 0/1/2 下的 AUC、PR-AUC、F1。F1 阈值仅从干净验证集选取，测试集不参与调参。"
        "补充了 K=1、K=3、K=5 三组原型数量对照实验。"
    )

    # 成果 1.1 主结果表
    doc.add_heading("成果 1.1　主结果表", level=2)
    add_table(doc,
        ["方法", "AUC（均值±标准差）", "PR-AUC（均值±标准差）", "F1（均值±标准差）"],
        [
            ["ResNet18 无增强",          "0.9849 ± 0.0032", "0.9636 ± 0.0037", "0.8922 ± 0.0182"],
            ["ResNet18 强增强",          "0.9963 ± 0.0012", "0.9881 ± 0.0024", "0.9237 ± 0.0107"],
            ["PBIP-Lite (β=0)",         "0.9950 ± 0.0001", "0.9865 ± 0.0012", "0.9357 ± 0.0212"],
            ["PBIP-Contrast (β=0.05)",  "0.9950 ± 0.0004", "0.9866 ± 0.0015", "0.9318 ± 0.0105"],
        ])
    doc.add_paragraph("数据来源：runs/summary_v3/main_results_summary.csv")

    # 成果 1.2 指标对比图
    doc.add_heading("成果 1.2　指标对比图", level=2)
    add_fig(doc, "paper_figs/main_metrics_comparison.png", "图 1　四种方法 AUC / PR-AUC / F1 指标对比柱状图")

    # 成果 1.3 K值结果
    doc.add_heading("成果 1.3　K 值结果", level=2)
    add_table(doc,
        ["K", "AUC", "PR-AUC", "F1", "Precision", "Recall"],
        [
            ["1", "0.9951", "0.9877", "0.9424", "0.9677", "0.9184"],
            ["3", "0.9954", "0.9883", "0.9430", "0.9579", "0.9286"],
            ["5", "0.9951", "0.9862", "0.9495", "0.9400", "0.9592"],
        ])
    doc.add_paragraph("数据来源：runs/ablations/k_ablation_results.csv")
    doc.add_paragraph("")
    add_fig(doc, "paper_figs/k_ablation_curve.png", "图 2　K=1/3/5 消融指标折线图")

    # 成果 1.4 方法结论
    doc.add_heading("成果 1.4　方法结论", level=2)
    doc.add_paragraph(
        "• 平均 AUC 最佳方法为 ResNet18 强增强：0.99625 ± 0.00125。\n"
        "• 平均 PR-AUC 最佳方法为 ResNet18 强增强：0.98814 ± 0.00244。\n"
        "• 平均 F1 最佳方法为 PBIP-Lite：0.93569 ± 0.02123。\n"
        "• K=5 时测试 F1 达到 0.9495，高于 K=1（0.9424）和 K=3（0.9430）。\n"
        "• 应同时讨论 PBIP 的平均性能与跨 seed 方差，不能只引用单个 seed。\n"
        "• 本表基于 1:3 负采样后的候选分类；FROC/CPM 不能表述为官方全候选 LUNA16 结果。"
    )
    doc.add_paragraph("完整结论模板：runs/summary_v3/conclusion_template.md")

    deliverable_list(doc, [
        "主结果表 → runs/summary_v3/main_results_summary.csv",
        "指标对比图 → paper_figs/main_metrics_comparison.png",
        "K 值结果 → runs/ablations/k_ablation_results.csv + paper_figs/k_ablation_curve.png",
        "方法结论 → runs/summary_v3/conclusion_template.md",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 2：消融结果表、配置文件、训练日志
    # ══════════════════════════════════════════
    doc.add_heading("任务 2　消融实验", level=1)
    doc.add_paragraph(
        "完成四组消融，每组独立配置 + 训练 + 评估：\n"
        "① 去掉原型 logits（仅保留对比损失辅助，α=0）\n"
        "② 去掉对比损失（仅保留原型融合 logits，β=0）\n"
        "③ 不同 K 值：K=1 / 3 / 5\n"
        "④ 不同 β 值：β=0.01 / 0.05 / 0.1 / 0.2"
    )

    # 成果 2.1 消融结果表
    doc.add_heading("成果 2.1　消融结果表", level=2)

    doc.add_paragraph("组件消融：", style="Intense Quote")
    add_table(doc,
        ["消融配置", "AUC", "PR-AUC", "F1"],
        [
            ["完整 PBIP-Contrast（基准）", "0.9954", "0.9883", "0.9430"],
            ["去掉原型 logits (α=0)",     "0.9947", "0.9874", "0.9375"],
            ["去掉对比损失 (β=0)",         "0.9950", "0.9876", "0.9179"],
        ])
    doc.add_paragraph("")

    doc.add_paragraph("β 值消融：", style="Intense Quote")
    add_table(doc,
        ["β", "验证 AUC", "测试 AUC", "测试 PR-AUC", "测试 F1"],
        [
            ["0.01", "0.9935", "0.9949", "0.9862", "0.9314"],
            ["0.05", "0.9914", "0.9954", "0.9883", "0.9430"],
            ["0.10", "0.9914", "0.9958", "0.9885", "0.9381"],
            ["0.20", "0.9930", "0.9946", "0.9863", "0.9307"],
        ])
    doc.add_paragraph("")
    add_fig(doc, "paper_figs/beta_ablation_curve.png", "图 3　β 值消融指标曲线")
    doc.add_paragraph("完整消融结果：runs/ablations/ablation_results.csv")

    # 成果 2.2 配置文件
    doc.add_heading("成果 2.2　配置文件", level=2)
    doc.add_paragraph("统一消融配置：src/configs/ablation.yaml，包含 K 值列表、β 值列表、组件开关和基准实验路径映射。")

    # 成果 2.3 训练日志
    doc.add_heading("成果 2.3　训练日志", level=2)
    doc.add_paragraph(
        "每组消融的完整训练日志保存在 runs/ablations/logs/ 下，包括：\n"
        "• no_prototype_logits.log — 去掉原型 logits 的 20 epoch 训练日志\n"
        "• k1.log / k5.log — K=1 和 K=5 的训练日志\n"
        "• build_bank_k1.log / build_bank_k5.log — 原型库构建日志\n"
        "• beta_0.01.log / beta_0.2.log — 新增 β 值训练日志\n"
        "• 每组配套 training_curve.csv 和 training_curves.png 训练曲线"
    )

    deliverable_list(doc, [
        "消融结果表 → runs/ablations/ablation_results.csv",
        "配置文件 → src/configs/ablation.yaml",
        "训练日志 → runs/ablations/logs/*.log + 各组 training_curve.csv",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 3：退化脚本、参数配置、样本网格图
    # ══════════════════════════════════════════
    doc.add_heading("任务 3　图像退化算子", level=1)
    doc.add_paragraph(
        "针对 CT patch 实现 6 种退化，均支持 1-5 档强度参数化。"
        "输入输出保持 C×H×W 和 [0,1] 范围，兼容 NumPy 和 Tensor。"
    )

    # 成果 3.1 退化脚本
    doc.add_heading("成果 3.1　退化脚本", level=2)
    add_table(doc,
        ["退化类型", "参数说明", "实现方式"],
        [
            ["低对比度",       "缩放因子 0.90→0.30",          "以每通道均值为中心线性缩放"],
            ["高斯噪声",       "σ = 0.01→0.10",              "逐像素加性高斯噪声"],
            ["高斯模糊",       "核大小 3→9，σ=0.5→2.0",     "OpenCV GaussianBlur"],
            ["窗宽窗位偏移",   "WW +100→+700；WL +25→+250", "由原肺窗反归一化后重加窗"],
            ["JPEG 压缩",      "质量因子 90→15",             "OpenCV JPEG 编解码"],
            ["下采样再上采样", "缩放比例 0.85→0.25",         "区域下采样 + 双线性上采样"],
        ])
    doc.add_paragraph("脚本位置：src/degradation.py（含 DegradationTransform 和 MixedDegradationTransform 两个类）")

    # 成果 3.2 参数配置
    doc.add_heading("成果 3.2　参数配置", level=2)
    doc.add_paragraph("配置位置：src/configs/degradation.yaml，定义了每种退化在 Level 1-5 下的具体参数值。")

    # 成果 3.3 样本网格图
    doc.add_heading("成果 3.3　样本网格图", level=2)
    add_fig(doc, "paper_figs/degradation_grid.png", "图 4　6 种退化类型 × N 个测试样本（Level 3）网格")

    deliverable_list(doc, [
        "退化脚本 → src/degradation.py",
        "参数配置 → src/configs/degradation.yaml",
        "样本网格图 → paper_figs/degradation_grid.png",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 4：鲁棒性结果表、退化曲线
    # ══════════════════════════════════════════
    doc.add_heading("任务 4　退化鲁棒性评估", level=1)
    doc.add_paragraph(
        "对 4 种方法分别在 5 档退化强度下测试 AUC / F1。"
        "每种方法从 seed 0/1/2 中按验证 AUC 选取代表模型；F1 阈值在干净验证集选择一次后固定。"
    )

    # 成果 4.1 鲁棒性结果表
    doc.add_heading("成果 4.1　鲁棒性结果表", level=2)
    add_table(doc,
        ["方法", "Level 0（干净）", "Level 1", "Level 3", "Level 5", "Level 5 F1 下降"],
        [
            ["ResNet18 基线",   "AUC=0.987 F1=0.897", "AUC=0.988 F1=0.893", "AUC=0.985 F1=0.863", "AUC=0.955 F1=0.739", "▼ 0.158"],
            ["ResNet18 强增强", "AUC=0.995 F1=0.914", "AUC=0.995 F1=0.916", "AUC=0.995 F1=0.912", "AUC=0.974 F1=0.846", "▼ 0.069"],
            ["PBIP-Lite",       "AUC=0.995 F1=0.930", "AUC=0.995 F1=0.928", "AUC=0.995 F1=0.929", "AUC=0.973 F1=0.865", "▼ 0.065"],
            ["PBIP-Contrast",   "AUC=0.995 F1=0.930", "AUC=0.996 F1=0.929", "AUC=0.995 F1=0.930", "AUC=0.973 F1=0.868", "▼ 0.062"],
        ])
    doc.add_paragraph("完整明细：runs/robustness/robustness_detailed.csv")
    doc.add_paragraph("汇总统计：runs/robustness/robustness_summary.csv")

    # 成果 4.2 退化曲线
    doc.add_heading("成果 4.2　退化曲线", level=2)
    add_fig(doc, "paper_figs/robustness_auc_curves.png", "图 5　退化强度 – AUC 下降曲线")
    add_fig(doc, "paper_figs/robustness_f1_curves.png", "图 6　退化强度 – F1 下降曲线")
    doc.add_paragraph(
        "观察结果：PBIP 系列方法在 Level 5 严重退化下 F1 下降 0.062-0.065，"
        "小于无增强 ResNet18 的 0.158，也略小于强增强 ResNet18 的 0.069。"
        "该结果支持 PBIP 在当前 sampled-candidate 测试集上具有较好的退化稳定性，"
        "但不能据此单独证明原型融合是性能差异的唯一原因。"
    )

    deliverable_list(doc, [
        "鲁棒性结果表 → runs/robustness/robustness_detailed.csv + robustness_summary.csv",
        "退化曲线 → paper_figs/robustness_auc_curves.png + robustness_f1_curves.png",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 5：最佳权重、训练日志、结果表
    # ══════════════════════════════════════════
    doc.add_heading("任务 5　Robust-PBIP 训练", level=1)
    doc.add_paragraph(
        "在 PBIP-Lite 基础上，训练集按 35% 概率保留干净样本、65% 概率随机施加 1-4 档退化。"
        "验证集始终保持干净。训练 20 epoch，最佳权重按验证 AUC 保存。"
    )

    # 成果 5.1 最佳权重
    doc.add_heading("成果 5.1　最佳权重", level=2)
    doc.add_paragraph("位置：runs/robust_pbip/seed_0/best_model.pth（最佳验证 AUC 出现在 Epoch 17，val_auc = 0.9928）")

    # 成果 5.2 训练日志
    doc.add_heading("成果 5.2　训练日志", level=2)
    doc.add_paragraph("训练曲线 CSV：runs/robust_pbip/seed_0/training_curve.csv")
    doc.add_paragraph("训练配置：runs/robust_pbip/seed_0/config.json")
    add_fig(doc, "runs/robust_pbip/seed_0/training_curves.png", "图 7　Robust-PBIP 训练曲线", width=Inches(5.0))

    # 成果 5.3 结果表
    doc.add_heading("成果 5.3　结果表（干净 vs 退化对比）", level=2)
    add_table(doc,
        ["模型", "场景", "AUC", "PR-AUC", "F1"],
        [
            ["PBIP",        "干净测试集",               "0.9954", "0.9883", "0.9430"],
            ["PBIP",        "退化 Level 3（六类均值）",  "0.9938", "0.9858", "0.9412"],
            ["Robust-PBIP", "干净测试集",               "0.9957", "0.9873", "0.8879"],
            ["Robust-PBIP", "退化 Level 3（六类均值）",  "0.9955", "0.9868", "0.8956"],
        ])
    doc.add_paragraph("完整对比：runs/robust_pbip/seed_0/clean_degraded_comparison.csv")
    doc.add_paragraph(
        "结果解释：Robust-PBIP 的退化宏平均 AUC 从普通 PBIP 的 0.9938 提升到 0.9955，"
        "但在验证集固定阈值下，退化宏平均 F1 从 0.9412 降至 0.8956。"
        "因此本实验只能表述为排序指标 AUC 的鲁棒性改善；F1 与概率校准仍需进一步优化。"
    )

    deliverable_list(doc, [
        "最佳权重 → runs/robust_pbip/seed_0/best_model.pth",
        "训练日志 → runs/robust_pbip/seed_0/training_curve.csv + training_curves.png",
        "结果表 → runs/robust_pbip/seed_0/clean_degraded_comparison.csv",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 6：Grad-CAM代码、热图、典型样例图
    # ══════════════════════════════════════════
    doc.add_heading("任务 6　Grad-CAM + 四类样本可视化", level=1)
    doc.add_paragraph(
        "为 ResNet18（强增强）和 PBIP-Contrast 分别实现了适配 3 通道 2.5D 输入的 Grad-CAM。"
        "目标层为 ResNet18 的 layer4 最后一个 BasicBlock。"
        "从测试集自动筛选 TP / FP / FN / TN 典型样本，每类最多 3 个。"
    )

    # 成果 6.1 Grad-CAM 代码
    doc.add_heading("成果 6.1　Grad-CAM 代码", level=2)
    doc.add_paragraph("位置：src/gradcam.py（含 GradCAM 类、自动四类样本筛选、热图叠加绘制功能）")

    # 成果 6.2 热图
    doc.add_heading("成果 6.2　热图", level=2)
    add_fig(doc, "paper_figs/gradcam_resnet18_augmented_four_categories.png",
            "图 8　ResNet18（强增强）TP/FP/FN/TN 四类 Grad-CAM 热力图")
    add_fig(doc, "paper_figs/gradcam_pbip_full_four_categories.png",
            "图 9　PBIP-Contrast TP/FP/FN/TN 四类 Grad-CAM 热力图")

    # 成果 6.3 典型样例图
    doc.add_heading("成果 6.3　典型样例筛选记录", level=2)
    doc.add_paragraph(
        "ResNet18 筛选了 11 个典型样本（TP 3、FP 3、FN 2、TN 3），"
        "PBIP 筛选了 12 个典型样本（TP 3、FP 3、FN 3、TN 3）。\n"
        "筛选记录：runs/gradcam/gradcam_samples.csv"
    )

    deliverable_list(doc, [
        "Grad-CAM 代码 → src/gradcam.py",
        "热图 → paper_figs/gradcam_resnet18_augmented_four_categories.png",
        "热图 → paper_figs/gradcam_pbip_full_four_categories.png",
        "典型样例图筛选记录 → runs/gradcam/gradcam_samples.csv",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 7：检索代码、特征索引、病例展示图
    # ══════════════════════════════════════════
    doc.add_heading("任务 7　Top-3 相似病例检索", level=1)
    doc.add_paragraph(
        "基于 PBIP-Contrast 的 backbone 提取 512 维特征向量，"
        "为训练集（3553 条）和测试集（497 条）构建序列化检索索引。"
        "查询时排除同一 CT（seriesuid）的样本，使用余弦相似度排序返回 Top-3。"
    )

    # 成果 7.1 检索代码
    doc.add_heading("成果 7.1　检索代码", level=2)
    doc.add_paragraph("位置：src/retrieval.py（含特征索引构建、Top-K 余弦检索、可视化展示功能）")

    # 成果 7.2 特征索引
    doc.add_heading("成果 7.2　特征索引", level=2)
    doc.add_paragraph(
        "训练集索引：runs/retrieval/pbip_full_train_index.npz（3553 条 × 512 维）\n"
        "测试集索引：runs/retrieval/pbip_full_test_index.npz（497 条 × 512 维）\n"
        "索引清单：runs/retrieval/pbip_full_index_manifest.json"
    )

    # 成果 7.3 病例展示图
    doc.add_heading("成果 7.3　病例展示图", level=2)
    add_table(doc,
        ["排名", "真实标签", "余弦相似度", "模型预测概率"],
        [
            ["Top-1", "阳性（结节）", "0.9995", "≈1.0000"],
            ["Top-2", "阳性（结节）", "0.9994", "≈1.0000"],
            ["Top-3", "阳性（结节）", "0.9994", "≈1.0000"],
        ])
    doc.add_paragraph("")
    add_fig(doc, "paper_figs/retrieval_top3_example.png", "图 10　Top-3 相似病例检索展示")

    deliverable_list(doc, [
        "检索代码 → src/retrieval.py",
        "特征索引 → runs/retrieval/pbip_full_train_index.npz + test_index.npz",
        "病例展示图 → paper_figs/retrieval_top3_example.png",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════
    # 任务 8：可行性报告、非重叠病例清单、处理记录
    # ══════════════════════════════════════════
    doc.add_heading("任务 8　LIDC-IDRI 外部验证可行性分析", level=1)
    doc.add_paragraph(
        "编写了独立的 LIDC-IDRI 数据处理脚本（src/lidc_external.py），"
        "使用 SimpleITK/GDCM 读取 DICOM、标准库解析 XML 标注，无需额外安装 pydicom。"
    )

    # 成果 8.1 可行性报告
    doc.add_heading("成果 8.1　可行性报告", level=2)
    doc.add_paragraph(
        "报告位置：runs/external_validation/feasibility_report.md\n\n"
        "当前状态：DATA_NOT_PRESENT（外部数据尚未下载）。\n"
        "全部代码逻辑已通过轻量单元测试验证。\n\n"
        "正式外部验证前检查清单：\n"
        "• 下载并校验完整 LIDC-IDRI DICOM 与 XML\n"
        "• 确认非重叠病例清单不含任何 LUNA16 series\n"
        "• 冻结阅片医师共识到二分类标签的规则\n"
        "• 生成不依赖测试标签的外部负候选\n"
        "• 冻结模型与阈值后一次性执行外部评估"
    )

    # 成果 8.2 非重叠病例清单
    doc.add_heading("成果 8.2　非重叠病例清单", level=2)
    doc.add_paragraph(
        "去重统计：runs/external_validation/deduplication_stats.json\n"
        "（LUNA16 共 888 个 seriesuid 用于精确匹配去重）\n"
        "非重叠病例清单：runs/external_validation/non_overlap_cases.csv\n"
        "由于当前未提供 LIDC-IDRI 数据，DICOM series、XML 和非重叠病例数均为 0；"
        "该 CSV 目前仅为结构化空清单，不能作为已完成真实去重的证据。"
    )

    # 成果 8.3 处理记录
    doc.add_heading("成果 8.3　处理记录", level=2)
    doc.add_paragraph(
        "DICOM/XML 处理日志：\n"
        "• runs/external_validation/xml_processing_log.csv\n"
        "• runs/external_validation/patch_processing_log.csv\n"
        "• runs/external_validation/lidc_case_inventory.csv\n"
        "• runs/external_validation/lidc_annotations.csv\n"
        "当前 patch_processing_log.csv 状态为 NOT_RUN；脚本、坐标逻辑和 XML 解析测试已完成，"
        "真实外部 patch 提取需在数据下载后执行。"
    )

    deliverable_list(doc, [
        "可行性报告 → runs/external_validation/feasibility_report.md",
        "非重叠病例清单 → runs/external_validation/non_overlap_cases.csv",
        "处理记录 → runs/external_validation/xml_processing_log.csv 等",
    ])

    # ═══ 保存 ═══
    doc.save(str(output_path))
    print(f"[OK] Word: {output_path}")


# ───────────────── 打包 ZIP ─────────────────

def package_zip(zip_path: Path, report_path: Path):
    files = [
        # 报告本身
        report_path.relative_to(PROJECT).as_posix(),

        # 任务 1：主结果表、指标对比图、K值结果、方法结论
        "runs/summary_v3/main_results.csv",
        "runs/summary_v3/main_results_summary.csv",
        "paper_figs/main_metrics_comparison.png",
        "runs/ablations/k_ablation_results.csv",
        "paper_figs/k_ablation_curve.png",
        "runs/summary_v3/conclusion_template.md",
        "runs/summary_v3/summary_manifest.json",

        # 任务 2：消融结果表、配置文件、训练日志
        "runs/ablations/ablation_results.csv",
        "runs/ablations/beta_ablation_results.csv",
        "runs/ablations/ablation_summary.json",
        "runs/ablations/ablation_manifest.json",
        "runs/ablations/ablation_conclusion_template.md",
        "src/configs/ablation.yaml",
        "paper_figs/beta_ablation_curve.png",
        "runs/ablations/k1/config.json",
        "runs/ablations/k1/results.json",
        "runs/ablations/k1/training_curve.csv",
        "runs/ablations/k1/training_curves.png",
        "runs/ablations/k5/config.json",
        "runs/ablations/k5/results.json",
        "runs/ablations/k5/training_curve.csv",
        "runs/ablations/k5/training_curves.png",
        "runs/ablations/no_prototype_logits/config.json",
        "runs/ablations/no_prototype_logits/results.json",
        "runs/ablations/no_prototype_logits/training_curve.csv",
        "runs/ablations/no_prototype_logits/training_curves.png",
        "runs/ablations/beta_0.01/config.json",
        "runs/ablations/beta_0.01/results.json",
        "runs/ablations/beta_0.01/training_curve.csv",
        "runs/ablations/beta_0.01/training_curves.png",
        "runs/ablations/beta_0.2/config.json",
        "runs/ablations/beta_0.2/results.json",
        "runs/ablations/beta_0.2/training_curve.csv",
        "runs/ablations/beta_0.2/training_curves.png",
        "runs/ablations/prototype_bank_k1/prototype_grid.png",
        "runs/ablations/prototype_bank_k1/cluster_stats.csv",
        "runs/ablations/prototype_bank_k1/config.json",
        "runs/ablations/prototype_bank_k5/prototype_grid.png",
        "runs/ablations/prototype_bank_k5/cluster_stats.csv",
        "runs/ablations/prototype_bank_k5/config.json",

        # 任务 3：退化脚本、参数配置、样本网格图
        "src/degradation.py",
        "src/configs/degradation.yaml",
        "paper_figs/degradation_grid.png",

        # 任务 4：鲁棒性结果表、退化曲线
        "runs/robustness/robustness_detailed.csv",
        "runs/robustness/robustness_summary.csv",
        "runs/robustness/selected_models.csv",
        "runs/robustness/evaluation_protocol.json",
        "paper_figs/robustness_auc_curves.png",
        "paper_figs/robustness_f1_curves.png",

        # 任务 5：最佳权重、训练日志、结果表
        "runs/robust_pbip/seed_0/best_model.pth",
        "runs/robust_pbip/seed_0/training_curve.csv",
        "runs/robust_pbip/seed_0/training_curves.png",
        "runs/robust_pbip/seed_0/results.json",
        "runs/robust_pbip/seed_0/config.json",
        "runs/robust_pbip/seed_0/clean_degraded_comparison.csv",
        "src/train_robust_pbip.py",
        "src/configs/robust_pbip.yaml",

        # 任务 6：Grad-CAM代码、热图、典型样例图
        "src/gradcam.py",
        "runs/gradcam/gradcam_samples.csv",
        "runs/gradcam/resnet18_augmented_samples.csv",
        "runs/gradcam/pbip_full_samples.csv",
        "paper_figs/gradcam_resnet18_augmented_four_categories.png",
        "paper_figs/gradcam_pbip_full_four_categories.png",

        # 任务 7：检索代码、特征索引、病例展示图
        "src/retrieval.py",
        "runs/retrieval/pbip_full_train_index.npz",
        "runs/retrieval/pbip_full_test_index.npz",
        "runs/retrieval/pbip_full_index_manifest.json",
        "runs/retrieval/retrieval_example.csv",
        "paper_figs/retrieval_top3_example.png",

        # 任务 8：可行性报告、非重叠病例清单、处理记录
        "src/lidc_external.py",
        "runs/external_validation/feasibility_report.md",
        "runs/external_validation/non_overlap_cases.csv",
        "runs/external_validation/deduplication_stats.json",
        "runs/external_validation/lidc_case_inventory.csv",
        "runs/external_validation/lidc_annotations.csv",
        "runs/external_validation/xml_processing_log.csv",
        "runs/external_validation/patch_processing_log.csv",

        # 辅助代码
        "src/evaluate_robustness.py",
        "src/summarize_results.py",
        "src/run_ablations.py",
        "src/parse_ablation_logs.py",
        "src/experiment_utils.py",
        "src/configs/robustness.yaml",
        "src/generate_polished_report.py",
        "README.md",
        "PROJECT_STATUS.md",
        "INCREMENTAL_TASKS.md",
        "requirements.txt",
        "src/luna16_dataset.py",
        "src/train.py",
        "src/pbip_train.py",
        "src/prototype_bank.py",
        "src/metrics.py",
        "tests/test_degradation.py",
        "tests/test_gradcam.py",
        "tests/test_retrieval.py",
        "tests/test_lidc_external.py",
        "tests/test_data_integrity.py",
        "tests/test_metrics.py",
        "tests/test_pbip.py",
        "tests/test_prototype_bank.py",
        "tests/test_run_experiments.py",
    ]

    dirs = ["runs/ablations/logs"]

    with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            full = PROJECT / f if not Path(f).is_absolute() else Path(f)
            if full.exists():
                zf.write(str(full), arcname=f)
            else:
                print(f"  [SKIP] {f}")
        for d in dirs:
            dp = PROJECT / d
            if dp.exists():
                for root, _, fnames in os.walk(dp):
                    for fn in fnames:
                        fp = Path(root) / fn
                        zf.write(str(fp), arcname=str(fp.relative_to(PROJECT)))

    size_mb = zip_path.stat().st_size / 1024 / 1024
    print(f"[OK] ZIP: {zip_path}  ({size_mb:.1f} MB)")


if __name__ == "__main__":
    out_dir = PROJECT / "deliverables"
    out_dir.mkdir(exist_ok=True)
    report = out_dir / "第三次增量实验成果报告.docx"
    zipf = out_dir / "第三次增量实验成果包.zip"
    create_report(report)
    package_zip(zipf, report)
    print("\n全部完成！")
